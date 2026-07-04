"""
=============================================================================
agwise_phenology_utils.py
=============================================================================
Shared helpers for the AgWise-style Sentinel-1/2 phenology pipeline.
Imported by all five pipeline scripts so that every script builds the
EXACT SAME paths and can find each other's outputs automatically:

  script1_Download_Stack_Smooth.py   → builds the multi-index Sentinel stack
  script1b_Variable_Importance_RF.py → ranks indices, writes best_index.txt
  script2_Threshold_Phenology.py     → threshold method
  script3_Derivative_Phenology.py    → 3 derivative methods
  script4_CrossValidate_Phenology.py → compares all methods vs ground truth

DIRECTORY CONVENTION ( AgWise R pipeline; MODIS → Sentinel):

  {base_dir}/agwise-datasourcing/dataops/datasourcing/Data/
    useCase_{country}_{useCaseName}/Sentinel/transform/MultiIndex/
      {country}_{useCaseName}_Sentinel_MultiIndex_{Py}_{Hy}_SG.tif
      {country}_{useCaseName}_Sentinel_MultiIndex_{Py}_{Hy}_band_metadata.csv

  {base_dir}/agwise-datacuration/dataops/datacuration/Data/
    useCase_{country}_{useCaseName}/{crop}/raw/
      data4RS_{year}.csv          (ground truth, one file per year; or a
                                    single data4RS.csv with a "year" column)

  {base_dir}/agwise-potentialyield/dataops/potentialyield/Data/
    useCase_{country}_{useCaseName}/{crop}/result/
      VariableImportance/
        best_index.txt
        importance_pooled.csv, importance_*_by_year.csv, *.png
      RSPlantingDate/
        {country}_{useCaseName}_{crop}_{Py}_{Hy}_{Method}_PlantingDate.tif
        ..._{Method}_HarvestDate.tif
        ..._{Method}_CropDuration.tif
        ..._{Method}_validation_*.png / map_*.png
        ..._CrossValidation_*.png / *.csv
=============================================================================
"""

import ssl, os, re, warnings, calendar
os.environ.setdefault("SHAPE_RESTORE_SHX", "YES")

# TLS verification stays ON by default. Some CGIAR/corporate networks sit
# behind a TLS-intercepting proxy that breaks certificate validation; ONLY
# there, set AGWISE_INSECURE_SSL=1 to fall back to an unverified context.
# Disabling it unconditionally for the whole process (as before) silently
# exposes GEE credentials and every download to interception.
if os.environ.get("AGWISE_INSECURE_SSL") == "1":
    ssl._create_default_https_context = ssl._create_unverified_context
    warnings.warn("TLS certificate verification DISABLED (AGWISE_INSECURE_SSL=1).")

from pathlib import Path
from datetime import datetime, timedelta

import numpy as np
import pandas as pd
import rasterio
from rasterio.warp import reproject
from rasterio.enums import Resampling
from rasterio.mask import mask as rio_mask
from rasterio.features import geometry_mask
from scipy.signal import savgol_filter
import matplotlib.pyplot as plt
import matplotlib.colors as mcolors
from scipy.stats import pearsonr
from sklearn.metrics import confusion_matrix, accuracy_score

# Silence ONLY the known-noisy, expected-by-design numpy warnings (all-NaN
# slices and invalid divides are routine here because cropland masking sets
# huge swaths of a province-scale AOI to NaN). A blanket ignore() also hid
# the pipeline-logic warnings — including the cross-year / bad-data ones —
# that should reach the user.
for _msg in ("All-NaN slice encountered", "Mean of empty slice",
             "invalid value encountered", "Degrees of freedom"):
    warnings.filterwarnings("ignore", message=f".*{_msg}.*", category=RuntimeWarning)

MONTH_NUM = {"January":1,"February":2,"March":3,"April":4,"May":5,"June":6,
             "July":7,"August":8,"September":9,"October":10,"November":11,
             "December":12}
MONTH_DAYS = {"January":31,"February":28,"March":31,"April":30,"May":31,
              "June":30,"July":31,"August":31,"September":30,"October":31,
              "November":30,"December":31}

S2_INDICES_DEFAULT = ["EVI","NDVI","NDRE","NDMI","LSWI","GCVI"]
S1_BANDS_DEFAULT   = ["VV","VH","VHVV","RVI"]


# =============================================================================
# 1. ADMINISTRATIVE BOUNDARY  (R: geodata::gadm())
# =============================================================================

def _as_list(admin_unit_name):
    """Normalize admin_unit_name to a list (accepts a bare string too)."""
    if admin_unit_name is None:
        return None
    if isinstance(admin_unit_name, str):
        return [admin_unit_name]
    return list(admin_unit_name)


def get_gadm_boundary(country, level, admin_unit_name=None, timeout=15):
    """
    Fetch administrative boundary via GADM (geodata.ucdavis.edu).

    timeout : int, seconds to wait for the GADM server before raising a
        clear ConnectionError instead of hanging indefinitely.
        Default 15 s - long enough for a slow connection, short enough
        that a blocked/offline network fails fast rather than freezing
        the whole script.

    Prefer get_boundary() with aoi_path instead - it uses a local file
    and never touches the network, making it immune to GADM outages and
    corporate firewalls.
    """
    try:
        import pygadm
    except ImportError:
        raise ImportError("pygadm is required for GADM boundary lookup.\n"
                          "Install with: pip install pygadm\n"
                          "Or pass aoi_path= to get_boundary() to use a local file instead.")

    import requests

    # Inject a timeout into the underlying requests session that pygadm uses.
    # pygadm creates its own CachedSession; monkey-patch the timeout via a
    # requests adapter so the first DNS/TCP attempt fails fast.
    _orig_send = requests.Session.send
    def _send_with_timeout(self, request, **kwargs):
        kwargs.setdefault("timeout", timeout)
        return _orig_send(self, request, **kwargs)
    requests.Session.send = _send_with_timeout

    try:
        admin_unit_name = _as_list(admin_unit_name)
        gdf = pygadm.Items(name=country, content_level=level)
    except Exception as e:
        raise ConnectionError(
            f"\n{'='*60}\n"
            f"Cannot reach the GADM server (geodata.ucdavis.edu).\n"
            f"This is a network issue - GADM may be blocked by your\n"
            f"firewall, offline, or unreachable from your location.\n\n"
            f"FIX: Pass your local shapefile/GeoJSON to get_boundary():\n"
            f"    aoi_path  = r\"D:\\Agwise\\Western_dist.shp\"\n"
            f"    district_col = \"NAME_3\"   # or whichever column holds\n"
            f"                               # sub-county / district names\n"
            f"Original error: {e}\n"
            f"{'='*60}"
        ) from e
    finally:
        requests.Session.send = _orig_send   # always restore

    if admin_unit_name is not None:
        if level == 0:
            raise ValueError("admin_unit_name is not null, level can't be "
                             "eq. to 0 and should be set between 1 and 3")
        name_col = f"NAME_{level}"
        if name_col not in gdf.columns:
            raise ValueError(f"Column {name_col} not found. "
                             f"Available: {gdf.columns.tolist()}")
        gdf = gdf[gdf[name_col].isin(admin_unit_name)]
        if gdf.empty:
            raise ValueError(f"No matching units for {admin_unit_name} "
                             f"at level {level} in {country}")
    return gdf.to_crs("EPSG:4326")


def admin_unit_col(gdf, level):
    """Return the column name holding unit names at this admin level."""
    if level == 0:
        for c in ["COUNTRY","NAME_0"]:
            if c in gdf.columns: return c
        return gdf.columns[0]
    col = f"NAME_{level}"
    return col if col in gdf.columns else gdf.columns[0]


def _guess_district_col(gdf):
    """Best-effort guess at which column holds admin-unit names in a
    user-supplied local file (no NAME_<level> convention to rely on)."""
    candidates = ["district","District","DISTRICT","NAME_1","NAME_2","NAME_3",
                  "NAME","name","admin","ADM_NAME","shapeName"]
    for c in candidates:
        if c in gdf.columns:
            return c
    # fall back to first non-geometry text column
    for c in gdf.columns:
        if c != "geometry" and gdf[c].dtype == object:
            return c
    return gdf.columns[0]


def get_boundary(country, level, admin_unit_name=None, aoi_path=None, district_col=None):
    """
    Get the AOI boundary.

    LOCAL FILE (always preferred - no network needed):
        Pass aoi_path= pointing to any shapefile / GeoJSON / GeoPackage.
        GADM is skipped entirely; country/level/admin_unit_name are only
        used for output path naming, not for fetching anything.

    GADM FALLBACK (only if aoi_path is None):
        Downloads from geodata.ucdavis.edu via pygadm with a 15-second
        timeout. If the server is unreachable (common behind corporate
        firewalls or when GADM is down), a clear error is raised that
        tells you exactly which parameters to add.

    Parameters
    ----------
    aoi_path : str or None
        Path to a local boundary file. STRONGLY recommended - avoids
        all network dependency.
    district_col : str or None
        Column holding district/sub-county names. If None and aoi_path
        is given, auto-guessed (tries 'NAME_3', 'district', etc.).
        Ignored when using GADM (determined by level).

    Returns
    -------
    (geopandas.GeoDataFrame, district_col_name)
    """
    import geopandas as gpd

    if aoi_path is not None:
        gdf = gpd.read_file(aoi_path).to_crs("EPSG:4326")
        col = district_col or _guess_district_col(gdf)
        print(f"  [Boundary] Loaded from local file: {aoi_path}")
        print(f"  [Boundary] District column: '{col}'  ({len(gdf)} polygons)")
        admin_unit_name_list = _as_list(admin_unit_name)
        if admin_unit_name_list is not None and col in gdf.columns:
            sub = gdf[gdf[col].isin(admin_unit_name_list)]
            if not sub.empty:
                gdf = sub
        return gdf, col

    # No local file - try GADM (may fail on restricted networks)
    print(f"  [Boundary] No aoi_path given - attempting GADM download ...")
    print(f"  [Boundary] TIP: pass aoi_path=r'path/to/your.shp' to avoid\n"
         f"             network dependency entirely.")
    gdf = get_gadm_boundary(country, level, admin_unit_name)
    col = admin_unit_col(gdf, level)
    return gdf, col


# =============================================================================
# 2. CROP / CROP-TYPE MASKS
# =============================================================================

def get_cropland_mask_gee(AOI, source="ESA/WorldCover/v200", cropland_class=40):
    import ee
    wc = ee.ImageCollection(source).first().select("Map")
    return wc.eq(cropland_class).clip(AOI)

def get_croptype_mask_gee(AOI, source, band, value):
    import ee
    if source is None:
        print("  ⚠ CropType=True but no crop_type_source given - skipping "
              "crop-type mask (CropMask alone still applied if set).")
        return None
    img = ee.Image(source).select(band)
    return img.eq(value).clip(AOI)


# =============================================================================
# 3. PATH BUILDERS  (single source of truth for every script)
# =============================================================================

def path_stack_dir(base_dir, country, useCaseName):
    return (Path(base_dir) / "agwise-datasourcing" / "dataops" / "datasourcing"
            / "Data" / f"useCase_{country}_{useCaseName}"
            / "Sentinel" / "transform" / "MultiIndex")

def path_stack_file(base_dir, country, useCaseName, Planting_year, Harvesting_year):
    d = path_stack_dir(base_dir, country, useCaseName)
    fname = (f"{country}_{useCaseName}_Sentinel_MultiIndex_"
             f"{Planting_year}_{Harvesting_year}_SG.tif")
    return d / fname

def path_stack_meta(base_dir, country, useCaseName, Planting_year, Harvesting_year):
    d = path_stack_dir(base_dir, country, useCaseName)
    fname = (f"{country}_{useCaseName}_Sentinel_MultiIndex_"
             f"{Planting_year}_{Harvesting_year}_band_metadata.csv")
    return d / fname

def path_gt_dir(base_dir, country, useCaseName, crop):
    return (Path(base_dir) / "agwise-datacuration" / "dataops" / "datacuration"
            / "Data" / f"useCase_{country}_{useCaseName}" / crop / "raw")

def path_result_dir(base_dir, country, useCaseName, crop):
    return (Path(base_dir) / "agwise-potentialyield" / "dataops" / "potentialyield"
            / "Data" / f"useCase_{country}_{useCaseName}" / crop / "result")

def path_varimp_dir(base_dir, country, useCaseName, crop):
    return path_result_dir(base_dir, country, useCaseName, crop) / "VariableImportance"

def path_rsplanting_dir(base_dir, country, useCaseName, crop):
    return path_result_dir(base_dir, country, useCaseName, crop) / "RSPlantingDate"

def name_tag(country, useCaseName, crop, Planting_year, Harvesting_year, index=None, extra=None):
    base = f"{country}_{useCaseName}_{crop}_{Planting_year}_{Harvesting_year}"
    # `index` (optional, default None for backward compatibility with any
    # caller that doesn't care) keeps a single-index run (e.g. "NDRE") and a
    # combine_indices run (e.g. "EVI-NDRE-VHVV") on the same
    # country/useCase/crop/years from writing to the SAME filenames and
    # silently overwriting each other's rasters/maps/validation outputs.
    if index:
        base += f"_{index}"
    # `extra` is a free-form suffix for any other parameter that changes the
    # actual phenology output and therefore needs its own filenames - e.g.
    # a thr_sos/thr_eos calibration sweep, where every trial uses the same
    # country/useCase/crop/years/index but a different thr.
    if extra:
        base += f"_{extra}"
    return base


def read_ground_truth_file(path):
    """
    Read a ground-truth file in either CSV or R's native RDS/RData format.

    R's RDS is a binary serialization format (NOT text/CSV) - pandas can't
    read it directly. RDS support requires: pip install pyreadr

    Returns a pandas DataFrame either way, with date-like numeric columns
    (R epoch-day integers, if pyreadr didn't resolve them to datetimes)
    defensively converted so downstream `to_doy()` parsing still works.
    """
    path = Path(path)
    suffix = path.suffix.lower()

    if suffix == ".csv":
        df = pd.read_csv(path)
    elif suffix in (".rds", ".rdata"):
        try:
            import pyreadr
        except ImportError:
            raise ImportError(
                "Reading .RDS/.RData ground truth files requires pyreadr.\n"
                "Install with: pip install pyreadr")
        result = pyreadr.read_r(str(path))
        # RDS holds exactly one object; RData can hold several - take the
        # first (and for RDS, only) one.
        df = next(iter(result.values()))
    else:
        raise ValueError(f"Unsupported ground truth file type '{suffix}': {path}")

    # Defensive: R Date columns sometimes arrive as raw numeric (days since
    # 1970-01-01) rather than parsed datetimes, depending on the RDS's
    # internal structure / R version used to save it.
    for col in df.columns:
        if "date" in col.lower() and pd.api.types.is_numeric_dtype(df[col]):
            try:
                df[col] = pd.to_datetime(df[col], unit="D", origin="1970-01-01")
            except Exception:
                pass

    return df


def find_ground_truth_files(base_dir, country, useCaseName, crop):
    """
    Auto-discover ground-truth files matching '*data4RS*' (.csv, .rds, or
    .RData) and parse a year from the filename. Returns {year: Path}.
    If a file has no year token but contains a 'year' column internally,
    it's still returned with key 'pooled' so callers can split it.
    """
    gt_dir = path_gt_dir(base_dir, country, useCaseName, crop)
    raw_matches = (sorted(gt_dir.glob("*data4RS*.csv")) +
                  sorted(gt_dir.glob("*data4RS*.rds")) +
                  sorted(gt_dir.glob("*data4RS*.RDS")) +
                  sorted(gt_dir.glob("*data4RS*.RData")) +
                  sorted(gt_dir.glob("*data4RS*.rdata")))
    # On Windows (case-insensitive filesystem), "*.rds" and "*.RDS" can
    # both match the SAME physical file - dedupe by resolved path so it
    # isn't processed twice.
    seen_resolved, files = set(), []
    for f in raw_matches:
        rp = f.resolve()
        if rp not in seen_resolved:
            seen_resolved.add(rp)
            files.append(f)
    out = {}
    for f in files:
        m = re.search(r"(19|20)\d{2}", f.stem)
        if m:
            out[int(m.group(0))] = f
        else:
            out.setdefault("pooled", []).append(f)
    return out, gt_dir


# =============================================================================
# 4. SEASON / DATE HELPERS
# =============================================================================

def season_bounds(Planting_year, Harvesting_year, Planting_month, Harvesting_month):
    s = datetime(Planting_year, MONTH_NUM[Planting_month], 1)
    # Last calendar day of the harvest month. calendar.monthrange handles
    # February 29 in leap years — MONTH_DAYS hard-codes 28 and would drop
    # the 29th when the season ends in a leap-year February.
    last_day = calendar.monthrange(Harvesting_year, MONTH_NUM[Harvesting_month])[1]
    e = datetime(Harvesting_year, MONTH_NUM[Harvesting_month], last_day)
    return s, e

def date_list(start, end, step):
    out=[]; d=start
    while d < end:
        out.append(d); d += timedelta(days=step)
    return out


# =============================================================================
# 5. GEE COMPOSITE BUILDERS
# =============================================================================

def _mask_s2(img):
    scl = img.select("SCL")
    m = (scl.neq(3).And(scl.neq(8)).And(scl.neq(9))
            .And(scl.neq(10)).And(scl.neq(11)))
    return img.updateMask(m).divide(10000).copyProperties(img, ["system:time_start"])

def _add_s2_indices(img):
    N=img.select("B8"); R=img.select("B4"); B=img.select("B2")
    G=img.select("B3"); RE1=img.select("B5")
    SWIR1=img.select("B11"); SWIR2=img.select("B12")
    return img.addBands([
        img.expression("2.5*((N-R)/(N+6*R-7.5*B+1))",
                       {"N":N,"R":R,"B":B}).rename("EVI"),
        N.subtract(R).divide(N.add(R)).rename("NDVI"),
        N.subtract(RE1).divide(N.add(RE1)).rename("NDRE"),
        N.subtract(SWIR1).divide(N.add(SWIR1)).rename("NDMI"),
        N.subtract(SWIR2).divide(N.add(SWIR2)).rename("LSWI"),
        N.divide(G).subtract(1).rename("GCVI"),
    ])

def _add_s1_bands(img):
    import ee
    vv_db=img.select("VV"); vh_db=img.select("VH")
    lin_vv=ee.Image(10).pow(vv_db.divide(10))
    lin_vh=ee.Image(10).pow(vh_db.divide(10))
    vhvv=vh_db.subtract(vv_db).rename("VHVV")
    rvi =lin_vh.multiply(4).divide(lin_vv.add(lin_vh)).rename("RVI")
    return img.addBands([vhvv, rvi])

def get_s2_composite(AOI, index, start_str, end_str, cloud_thresh=80):
    import ee
    col = (ee.ImageCollection("COPERNICUS/S2_SR_HARMONIZED")
             .filterBounds(AOI).filterDate(start_str, end_str)
             .filter(ee.Filter.lt("CLOUDY_PIXEL_PERCENTAGE", cloud_thresh))
             .map(_mask_s2).map(_add_s2_indices).select(index))
    fb = ee.Image.constant(0).rename(index).toFloat()
    return ee.Image(ee.Algorithms.If(col.size().gt(0), col.median(), fb)
            ).clip(AOI).set("system:time_start", ee.Date(start_str).millis())

def get_s1_composite(AOI, band, start_str, end_str, orbit="DESCENDING"):
    import ee
    col = (ee.ImageCollection("COPERNICUS/S1_GRD")
             .filterBounds(AOI).filterDate(start_str, end_str)
             .filter(ee.Filter.listContains("transmitterReceiverPolarisation","VV"))
             .filter(ee.Filter.listContains("transmitterReceiverPolarisation","VH"))
             .filter(ee.Filter.eq("instrumentMode","IW"))
             .filter(ee.Filter.eq("orbitProperties_pass", orbit))
             .select(["VV","VH"]).map(_add_s1_bands).select(band))
    fb_val = 0 if band == "RVI" else -20
    fb = ee.Image.constant(fb_val).rename(band).toFloat()
    return ee.Image(ee.Algorithms.If(col.size().gt(0), col.median(), fb)
            ).clip(AOI).set("system:time_start", ee.Date(start_str).millis())


def download_composite(img, fpath, AOI, scale, overwrite=False, num_threads=4,
                       max_retries=5, retry_wait=10):
    """
    Download a single-band composite to disk, with:
      - caching (skip if already downloaded)
      - corrupt/partial-file detection (a crash mid-download can leave a
        broken .tif behind - this is re-downloaded rather than trusted)
      - automatic retry with exponential backoff on GEE's 429
        "Too Many Requests" / concurrency-limit error

    Raises RuntimeError after max_retries if the rate limit persists.
    """
    import geemap, time

    if fpath.exists() and not overwrite:
        try:
            with rasterio.open(fpath) as src:
                _ = src.read(1, window=((0, 1), (0, 1)))   # cheap validity check
            return fpath
        except Exception:
            print(f"    ⚠ Cached file looks corrupt, re-downloading: {fpath.name}")
            try: fpath.unlink()
            except Exception: pass

    last_exc = None
    for attempt in range(1, max_retries + 1):
        try:
            geemap.download_ee_image(image=img, filename=str(fpath), region=AOI,
                                     scale=scale, crs="EPSG:4326",
                                     num_threads=num_threads, overwrite=True)
            return fpath
        except Exception as e:
            last_exc = e
            msg = str(e)
            transient = (
                "429" in msg or "Too Many Requests" in msg or "concurrency limit" in msg
                or "Internal Server Error" in msg or "Bad Gateway" in msg
                or "Service Unavailable" in msg or "Gateway Time-out" in msg
                or any(code in msg for code in (" 500", " 502", " 503", " 504"))
            )
            if transient:
                wait = retry_wait * attempt
                print(f"    ⚠ Transient GEE error (attempt {attempt}/{max_retries}). "
                     f"Waiting {wait}s ...")
                time.sleep(wait)
                continue
            raise

    raise RuntimeError(f"Failed after {max_retries} retries: {fpath.name}\n"
                       f"Last error: {last_exc}")


# =============================================================================
# 6. RESAMPLING (S1 20m → S2 10m reference grid)
# =============================================================================

def load_band_resampled(tif_path, ref_transform=None, ref_crs=None, ref_h=None, ref_w=None):
    with rasterio.open(tif_path) as src:
        arr = src.read(1).astype(np.float32)
        nd  = src.nodata
        if nd is not None: arr[arr==nd] = np.nan
        arr[np.abs(arr) > 1e6] = np.nan

        if ref_transform is None:
            return arr, src.transform, src.crs, src.height, src.width

        if (src.transform == ref_transform and src.crs == ref_crs and
            src.height == ref_h and src.width == ref_w):
            return arr, ref_transform, ref_crs, ref_h, ref_w

        out = np.full((ref_h, ref_w), np.nan, dtype=np.float32)
        reproject(source=arr, destination=out,
                  src_transform=src.transform, src_crs=src.crs, src_nodata=np.nan,
                  dst_transform=ref_transform, dst_crs=ref_crs, dst_nodata=np.nan,
                  resampling=Resampling.bilinear)
        return out, ref_transform, ref_crs, ref_h, ref_w


# =============================================================================
# 6b. CROP MASK  (R: Section 2.3.2, `stacked_SG_s <- stacked_SG_s*cropmask`)
# =============================================================================

def load_crop_mask_resampled(mask_path, ref_transform, ref_crs, ref_h, ref_w,
                             mask_value=1):
    """
    Load a local crop-mask/crop-type GeoTIFF and align it to the stack's
    reference grid, the same way `load_band_resampled()` aligns an S1 band
    onto the S2 grid.

    mask_path : path to a classified raster (e.g. a cropland mask where
        1=crop, or a multi-class crop-type map like an RF Ensemble_Prediction
        raster where each class has its own integer code).
    mask_value : the pixel value in the raster that means "this is the crop
        of interest" (default 1 - a binary cropland mask). Set this to the
        class code for the crop of interest if using a multi-class map (e.g.
        mask_value=2 if maize is coded as 2 in your classification).

    Resampling uses NEAREST, not bilinear - averaging a categorical mask
    across pixel boundaries would invent fractional, meaningless class
    values right where it matters most (the field edges).

    Returns a boolean array (ref_h, ref_w): True = inside the crop mask.
    """
    with rasterio.open(mask_path) as src:
        arr = src.read(1).astype(np.float32)
        nd = src.nodata
        if nd is not None:
            arr[arr == nd] = np.nan

        if (src.transform == ref_transform and src.crs == ref_crs and
            src.height == ref_h and src.width == ref_w):
            aligned = arr
        else:
            aligned = np.full((ref_h, ref_w), np.nan, dtype=np.float32)
            reproject(source=arr, destination=aligned,
                      src_transform=src.transform, src_crs=src.crs, src_nodata=np.nan,
                      dst_transform=ref_transform, dst_crs=ref_crs, dst_nodata=np.nan,
                      resampling=Resampling.nearest)

    return aligned == mask_value


def apply_crop_mask(stack, mask_bool):
    """
    Set every pixel outside the crop mask to NaN, across every time step -
    R's equivalent of multiplying the whole stack by a 1/NA cropland raster
    before any phenology metric is computed, so peak/min/amplitude/threshold
    are all derived only from in-mask (e.g. maize) pixels and the final
    sowing/harvest/duration maps come out NaN everywhere else.
    """
    out = stack.copy()
    out[:, ~mask_bool] = np.nan
    return out


# =============================================================================
# 7. VECTORIZED SG SMOOTHING  (fast - no per-pixel Python loop)
# =============================================================================

def _interp_nan_chunk(chunk, doys):
    """
    chunk : (T, C) - a slice of pixels, possibly containing NaN
    doys  : (T,) ascending day-of-year values for each row

    Interior NaNs are filled via DOY-aware linear interpolation (using the
    REAL gap sizes between composites, not assuming equal spacing).
    Leading/trailing NaNs are filled with the nearest valid value (no
    extrapolation). Pure NumPy, fully vectorized over C - no pandas, no
    per-pixel Python loop, so this scales to hundreds of millions of pixels.
    """
    if not np.isnan(chunk).any():
        return chunk   # fast path - the common case once downloads succeed

    T, C = chunk.shape
    valid = ~np.isnan(chunk)
    row_idx = np.arange(T)[:, None]

    prev_idx = np.where(valid, row_idx, -1)
    prev_idx = np.maximum.accumulate(prev_idx, axis=0)

    rev_valid = valid[::-1]
    rev_idx = np.where(rev_valid, row_idx, -1)
    rev_idx = np.maximum.accumulate(rev_idx, axis=0)
    next_idx = np.where(rev_idx[::-1] < 0, -1, (T - 1) - rev_idx[::-1])

    has_prev, has_next = prev_idx >= 0, next_idx >= 0
    safe_prev = np.where(has_prev, prev_idx, 0)
    safe_next = np.where(has_next, next_idx, 0)

    val_prev = np.take_along_axis(chunk, safe_prev, axis=0)
    val_next = np.take_along_axis(chunk, safe_next, axis=0)
    doy_prev = doys[safe_prev]
    doy_next = doys[safe_next]
    doy_cur  = np.broadcast_to(doys[:, None], (T, C))

    denom = doy_next - doy_prev
    frac = np.where(denom != 0, (doy_cur - doy_prev) / np.where(denom == 0, 1, denom), 0.0)
    interp_val = val_prev + frac * (val_next - val_prev)

    fill_val = np.where(has_prev & has_next, interp_val,
              np.where(has_prev, val_prev,
              np.where(has_next, val_next, np.nan)))
    return np.where(valid, chunk, fill_val)


def _build_weight_matrix(doys, d_reg):
    T, T_reg = len(doys), len(d_reg)
    Wm = np.zeros((T_reg, T), dtype=np.float32)
    for i, d in enumerate(d_reg):
        if d <= doys[0]:
            Wm[i, 0] = 1.0
        elif d >= doys[-1]:
            Wm[i, -1] = 1.0
        else:
            j = max(0, min(np.searchsorted(doys, d) - 1, T - 2))
            d0, d1 = doys[j], doys[j+1]
            frac = (d - d0) / (d1 - d0) if d1 != d0 else 0.0
            Wm[i, j]   = 1 - frac
            Wm[i, j+1] = frac
    return Wm


def vectorized_regrid_and_smooth(raw_stack, doys, sg_window=7, sg_poly=3,
                                  regular_step=8, chunk_size=5_000_000):
    """
    raw_stack : (T,H,W), possibly containing NaN (cloud/missing composites)
    doys      : (T,) ascending day-of-year values, same for every pixel

    Processes pixels in chunks of `chunk_size` so peak memory stays bounded
    regardless of AOI size (no pandas - a DataFrame with H*W columns is
    catastrophically slow/memory-heavy for province-scale AOIs).

    1. Fill NaN along time axis (DOY-aware linear interp; nearest-value at
       edges) - vectorized per chunk, with a fast path when a chunk has no
       NaNs at all (the common case).
    2. Regrid to a regular `regular_step`-day grid via a precomputed linear
       interpolation weight matrix (single matrix multiply per chunk).
    3. Apply Savitzky-Golay filter vectorized along the time axis.

    Returns (smoothed_stack (T_reg,H,W), d_reg (T_reg,))
    """
    # This routine (regular grid, weight matrix, edge fills) assumes a
    # strictly ascending day-of-year axis. A season crossing the New Year
    # arrives as [...360, 361, 5, 12, ...] and would be SILENTLY mis-gridded
    # (Jan sorted before Dec, interpolated across a negative gap). Fail loud
    # instead of producing plausible-but-wrong phenology.
    doys = np.asarray(doys, dtype=float)
    if doys.size > 1 and not np.all(np.diff(doys) > 0):
        raise ValueError(
            "vectorized_regrid_and_smooth needs a strictly ascending day-of-year "
            "axis but got a non-monotonic one. This happens when the season "
            "crosses a calendar year (e.g. Sep->Mar): day-of-year wraps 365->1. "
            "Cross-year seasons are NOT supported via the DOY convention — keep "
            "planting and harvest within one calendar year, or use a date-based "
            "(days-since-planting) axis."
        )
    T, H, W = raw_stack.shape
    flat = raw_stack.reshape(T, -1)
    N = flat.shape[1]

    d_reg = np.arange(doys.min(), doys.max() + 1, regular_step, dtype=float)
    T_reg = len(d_reg)
    Wm = _build_weight_matrix(doys, d_reg)

    out = np.empty((T_reg, N), dtype=np.float32)
    win = min(sg_window, T_reg)
    win = win if win % 2 == 1 else win - 1
    do_sg = win >= sg_poly + 2

    for start in range(0, N, chunk_size):
        end = min(start + chunk_size, N)
        chunk = flat[:, start:end]
        filled = _interp_nan_chunk(chunk, doys)
        regrid = Wm @ filled
        if do_sg:
            regrid = savgol_filter(regrid, win, sg_poly, axis=0)
        out[:, start:end] = regrid

    return out.reshape(T_reg, H, W), d_reg


# =============================================================================
# 8. BAND-NAME PARSING  (e.g. "EVI_DOY152_SG" → ("EVI", 152))
# =============================================================================

def parse_band_name(bname):
    """("EVI_DOY152_SG" → ("EVI", 152)) and, for the date-labelled bands
    that cross-year stacks use ("EVI_20200914_SG"), the date's day-of-year
    ("EVI_20200914_SG" → ("EVI", 258)). Same (index, doy) contract."""
    parts = str(bname).split("_")
    try:
        doy_i = next(i for i,p in enumerate(parts) if p.startswith("DOY"))
        index = "_".join(parts[:doy_i])
        doy   = int(parts[doy_i].replace("DOY",""))
        return index, doy
    except StopIteration:
        for i, p in enumerate(parts):
            if len(p) == 8 and p.isdigit():
                try:
                    dt = datetime.strptime(p, "%Y%m%d")
                except ValueError:
                    continue
                return "_".join(parts[:i]), int(dt.timetuple().tm_yday)
        return None, None

def load_rf_importance_weights(base_dir, country, useCaseName, crop, indices):
    """
    Auto-read RF importance scores from Variable_Importance_RF's
    importance_pooled.csv and return a weight dict for the requested indices,
    normalised so they sum to 1 (so the weighted mean stays on the [0,1] scale
    after per-pixel min-max normalisation).

    Falls back to equal weights (with a warning) if:
      - the file doesn't exist yet (Variable_Importance_RF hasn't been run), or
      - any of the requested indices is missing from the file.

    Typical use:
        weights = load_rf_importance_weights(base_dir, country, useCaseName,
                                             crop, ["EVI","NDRE","VHVV"])
        stack, doys = combine_indices_pixelwise(full_stack, descs,
                                               ["EVI","NDRE","VHVV"],
                                               weights=weights)
    """
    csv_path = path_varimp_dir(base_dir, country, useCaseName, crop) / "importance_pooled.csv"
    if not csv_path.exists():
        print(f"  [Weights] importance_pooled.csv not found at {csv_path} - "
             f"using equal weights. Run Variable_Importance_RF first to "
             f"get data-driven weights.")
        return {idx: 1.0/len(indices) for idx in indices}

    df = pd.read_csv(csv_path)
    # Column names produced by Variable_Importance_RF - typically "index"
    # and "importance" (or "MeanDecreaseGini", "mean_importance", etc.).
    # Try the most common variants.
    name_col = next((c for c in df.columns if c.lower() in ("index","variable","feature","band")), None)
    imp_col  = next((c for c in df.columns if any(k in c.lower() for k in
                     ("importance","gini","gain","weight","score"))), None)
    if name_col is None or imp_col is None:
        print(f"  [Weights] Can't parse importance_pooled.csv columns {list(df.columns)} - "
             f"using equal weights.")
        return {idx: 1.0/len(indices) for idx in indices}

    df = df[[name_col, imp_col]].rename(columns={name_col:"index", imp_col:"importance"})
    df["index"] = df["index"].str.strip()
    missing = [idx for idx in indices if idx not in df["index"].values]
    if missing:
        print(f"  [Weights] {missing} not found in importance_pooled.csv - "
             f"using equal weights for all indices.")
        return {idx: 1.0/len(indices) for idx in indices}

    imp = {row["index"]: row["importance"] for _, row in df.iterrows()}
    raw = {idx: float(imp[idx]) for idx in indices}
    total = sum(raw.values())
    if total == 0:
        return {idx: 1.0/len(indices) for idx in indices}
    normed = {idx: v/total for idx, v in raw.items()}
    print(f"  [Weights] Loaded from importance_pooled.csv: "
         + ", ".join(f"{k}={v:.3f}" for k,v in normed.items()))
    return normed


def combine_indices_pixelwise(stack, descs, indices, weights=None,
                               combine_method="weighted_mean",
                               chunk_size=5_000_000, pca_sample=200_000):
    """
    Combine multiple indices/bands into ONE pixel-wise time series.

    Pipeline (same for both methods):
      1. Extract each index's (T_idx,H,W) sub-stack + DOY grid.
      2. Interpolate every index onto the UNION of all DOY grids (linear
         weight-matrix, same as the main smoothing step).
      3. Normalise per-pixel, per-index (min-max over time) so indices on
         very different native scales - EVI ~0-1 vs a dB SAR band often
         negative - contribute comparably regardless of raw magnitude.

    Then, depending on combine_method:

    "weighted_mean" (default):
        Step 4 - weighted average across indices (NaN-aware).
        `weights` can be a dict {index: weight} (e.g. RF importance scores
        from load_rf_importance_weights()) or None for equal weights.
        Simple, interpretable, and behaves predictably when one index has
        many NaN pixels.

    "pca":
        Step 4 - PC1 of the normalized, interpolated n-index stack,
        fit on a random sample of valid (pixel×time) observations and
        then applied chunk-wise to the full AOI.
        `weights` is ignored.
        More powerful when indices carry correlated but partially redundant
        information: PCA finds the linear combination that captures the most
        variance across ALL indices simultaneously, without requiring you
        to specify or guess importance weights. Typical use: if EVI and
        NDRE both track canopy greenup but with slightly different sensitivities,
        PCA automatically down-weights whichever is more redundant. Loadings
        are printed so you can see which index PCA weighted most.
        Caveat: PCA maximises variance, not agronomic relevance. If one index
        has anomalously high variance (e.g. a noisy SAR band), PC1 may lean
        toward it even if it's the least informative about phenology. Check
        the loadings printout; if they look unintuitive, fall back to
        weighted_mean with RF importance weights instead.

    Memory: all methods process pixels in chunks - peak memory scales with
    chunk_size regardless of AOI size, matching vectorized_regrid_and_smooth.

    Returns (combined_stack (T_common,H,W), doy_common (T_common,))
    """
    # ── Step 1-2: extract + interpolate onto shared DOY grid ─────────────────
    sub_stacks, doy_grids = {}, {}
    for idx in indices:
        sub, doys = select_index_bands(stack, descs, idx)
        sub_stacks[idx] = sub
        doy_grids[idx] = doys

    doy_common = np.array(sorted(set().union(*[set(d) for d in doy_grids.values()])),
                          dtype=float)
    Tc = len(doy_common)
    H, W = next(iter(sub_stacks.values())).shape[1:]
    N = H * W

    flats, weight_mats = {}, {}
    for idx in indices:
        doys = doy_grids[idx]
        flats[idx] = sub_stacks[idx].reshape(len(doys), N)
        weight_mats[idx] = (None if (len(doys) == Tc and np.allclose(doys, doy_common))
                            else _build_weight_matrix(doys, doy_common))

    # ── Helper: normalize one aligned chunk in-place ──────────────────────────
    def _norm_chunk(aligned):
        """Min-max normalize over time axis (axis=0), in-place."""
        pmin = np.nanmin(aligned, axis=0)
        pmax = np.nanmax(aligned, axis=0)
        rng_safe = np.where((pmax - pmin) == 0, np.float32(1.0), pmax - pmin)
        valid = ~np.isnan(aligned)
        np.subtract(aligned, pmin, out=aligned)
        np.divide(aligned, rng_safe, out=aligned)
        return aligned, valid

    # ────────────────────────────────────────────────────────────────────────
    # METHOD A: weighted mean
    # ────────────────────────────────────────────────────────────────────────
    if combine_method == "weighted_mean":
        if weights is None:
            weights = {idx: 1.0 for idx in indices}
        out = np.empty((Tc, N), dtype=np.float32)
        for start in range(0, N, chunk_size):
            end = min(start + chunk_size, N)
            sum_w  = np.zeros((Tc, end-start), dtype=np.float32)
            sum_wx = np.zeros((Tc, end-start), dtype=np.float32)
            for idx in indices:
                chunk = flats[idx][:, start:end]
                Wm = weight_mats[idx]
                aligned = chunk.astype(np.float32, copy=True) if Wm is None else (Wm @ chunk)
                aligned, valid = _norm_chunk(aligned)
                w = np.float32(weights.get(idx, 1.0))
                np.place(aligned, ~valid, np.float32(0.0))
                np.multiply(aligned, w, out=aligned)
                sum_wx += aligned
                sum_w[valid] += w
            nz = sum_w > 0
            out[:, start:end] = np.where(nz, sum_wx / np.where(nz, sum_w, np.float32(1.0)), np.nan)
        return out.reshape(Tc, H, W), doy_common

    # ────────────────────────────────────────────────────────────────────────
    # METHOD B: PCA - fit on sample, apply chunk-wise
    # ────────────────────────────────────────────────────────────────────────
    elif combine_method == "pca":
        from sklearn.decomposition import PCA

        n_idx = len(indices)
        print(f"  [PCA] Fitting PC1 on a sample of ≤{pca_sample:,} valid observations ...")

        # Sample a random subset of pixel×time pairs that are valid (non-NaN)
        # across ALL indices simultaneously - PCA requires complete rows.
        # We sample columns (pixels) from the first chunk to stay memory-light,
        # then collect rows that have no NaN in any index.
        rng = np.random.default_rng(0)
        sample_cols = rng.choice(N, size=min(N, pca_sample // Tc + 1), replace=False)
        sample_cols = np.sort(sample_cols)
        normed_sample = {}
        for idx in indices:
            chunk = flats[idx][:, sample_cols]
            Wm = weight_mats[idx]
            aligned = chunk.astype(np.float32, copy=True) if Wm is None else (Wm @ chunk)
            aligned, _ = _norm_chunk(aligned)
            normed_sample[idx] = aligned   # (Tc, n_sample_cols)

        # Build (n_obs, n_idx) matrix where n_obs = Tc × n_sample_cols
        # Each row is one (timestep, pixel) observation; each column is one index.
        X_list = []
        for idx in indices:
            X_list.append(normed_sample[idx].reshape(-1))   # (Tc*n_sample_cols,)
        X = np.stack(X_list, axis=1).astype(np.float32)    # (Tc*n_sample_cols, n_idx)

        # Drop rows where ANY index is NaN (PCA requires complete rows)
        valid_rows = ~np.any(np.isnan(X), axis=1)
        X_clean = X[valid_rows]
        print(f"  [PCA] {X_clean.shape[0]:,} complete observations (of {X.shape[0]:,} sampled)")
        if X_clean.shape[0] < n_idx:
            raise RuntimeError(
                f"PCA needs at least {n_idx} complete rows but only got "
                f"{X_clean.shape[0]}. Too many NaN pixels in the sample - "
                f"try a larger pca_sample or use combine_method='weighted_mean'.")

        pca = PCA(n_components=1)
        pca.fit(X_clean)
        loadings = dict(zip(indices, pca.components_[0]))
        var_explained = pca.explained_variance_ratio_[0] * 100
        print(f"  [PCA] PC1 explains {var_explained:.1f}% of variance in the sample")
        print(f"  [PCA] Loadings: " + ", ".join(f"{k}={v:+.3f}" for k,v in loadings.items()))
        # Sign convention: ensure PC1 is positively correlated with the first
        # index so the combined curve rises during the growing season rather than
        # inverting. If the first loading is negative, flip all loadings.
        if pca.components_[0, 0] < 0:
            pca.components_ *= -1
            print(f"  [PCA] Flipped sign so PC1 rises with {indices[0]} (growing season convention)")

        # Apply PC1 projection chunk-wise
        out = np.empty((Tc, N), dtype=np.float32)
        for start in range(0, N, chunk_size):
            end = min(start + chunk_size, N)
            # Build (Tc*(end-start), n_idx) block for this chunk
            cols = []
            nan_mask = np.zeros((Tc, end-start), dtype=bool)
            for idx in indices:
                chunk = flats[idx][:, start:end]
                Wm = weight_mats[idx]
                aligned = chunk.astype(np.float32, copy=True) if Wm is None else (Wm @ chunk)
                aligned, valid = _norm_chunk(aligned)
                nan_mask |= ~valid
                cols.append(aligned.reshape(-1))
            X_chunk = np.stack(cols, axis=1)        # (Tc*(end-start), n_idx)
            row_valid = ~np.any(np.isnan(X_chunk), axis=1)
            # Project valid rows; set NaN rows to NaN in output
            proj = np.full(X_chunk.shape[0], np.nan, dtype=np.float32)
            if row_valid.any():
                proj[row_valid] = pca.transform(X_chunk[row_valid])[:, 0].astype(np.float32)
            out[:, start:end] = proj.reshape(Tc, end-start)

        # Re-normalize the PC1 projection to [0,1] per pixel so it stays on
        # the same scale as the weighted-mean output and threshold values.
        out_H_W = out.reshape(Tc, H, W)
        p_min = np.nanmin(out_H_W, axis=0)
        p_max = np.nanmax(out_H_W, axis=0)
        rng_safe = np.where((p_max - p_min) == 0, np.float32(1.0), p_max - p_min)
        out_H_W = (out_H_W - p_min) / rng_safe
        return out_H_W.astype(np.float32), doy_common

    else:
        raise ValueError(f"combine_method must be 'weighted_mean' or 'pca', got '{combine_method}'")


def select_index_bands(stack, descs, index):
    """Return (sub_stack (T,H,W), doys (T,)) for one index, sorted by DOY."""
    sel = []
    for i, d in enumerate(descs):
        idx, doy = parse_band_name(d)
        if idx == index:
            sel.append((i, doy))
    if not sel:
        raise ValueError(f"No bands found for index '{index}'. "
                         f"Available: {sorted(set(parse_band_name(d)[0] for d in descs))}")
    sel.sort(key=lambda x: x[1])
    band_idx = [s[0] for s in sel]
    doys     = np.array([s[1] for s in sel], dtype=float)
    return stack[band_idx], doys


# =============================================================================
# 9. OUTLIER REPLACEMENT  (mean ± 1.5 SD → mean)
# =============================================================================

def replace_outliers(arr, factor=1.5):
    mn = np.nanmean(arr); sd = np.nanstd(arr) * factor
    out = arr.copy()
    out[(arr < mn - sd) | (arr > mn + sd)] = mn
    return out


# =============================================================================
# 10. PHENOLOGY METHODS - PIXEL-WISE, VECTORIZED
# =============================================================================

def _peak_and_limbs(stack, doys):
    peak_val = np.nanmax(stack, axis=0)   # NaN for all-NaN pixels - safe, no crash
    # nanargmax CRASHES on an all-NaN slice (unlike nanmax/nanmin, which just
    # warn and return NaN) - extremely common once a cropland mask sets huge
    # swaths of a province-scale AOI to NaN across every timestep. Substitute
    # a sentinel so argmax always has something to pick, then mask the
    # result back to NaN using peak_val (already correctly NaN there).
    safe_stack = np.where(np.isnan(stack), -np.inf, stack)
    # R's `max.pheno[stack[[i]] == peakmx.max] <- doy` loop has no guard
    # against re-assignment, so on a tie it keeps the LAST (latest-DOY)
    # matching image. np.argmax alone would keep the first; running it on
    # the time-reversed stack and mapping the index back reproduces R's
    # last-match-wins tie-break instead.
    T = safe_stack.shape[0]
    peak_idx = T - 1 - np.argmax(safe_stack[::-1], axis=0)
    peak_doy_rast = np.where(np.isnan(peak_val), np.nan, doys[peak_idx].astype(float))
    median_peak = float(np.nanmedian(peak_doy_rast))
    left_mask  = doys <= median_peak
    right_mask = doys >= median_peak
    return peak_val, median_peak, left_mask, right_mask

def threshold_phenology_pixelwise(stack, doys, thr_sos, thr_eos, emergence):
   
    peak_val, median_peak, left_mask, right_mask = _peak_and_limbs(stack, doys)

    stack_left  = np.where(left_mask[:,None,None],  stack, np.nan)
    stack_right = np.where(right_mask[:,None,None], stack, np.nan)
    min_left  = np.nanmin(stack_left,  axis=0)
    min_right = np.nanmin(stack_right, axis=0)

    amplitude    = peak_val - (min_left + min_right) / 2.0
    thresh_left  = min_left  + thr_sos * amplitude
    thresh_right = min_right + thr_eos * amplitude

    H, W = peak_val.shape
    sos_doy = np.full((H,W), np.nan, dtype=np.float32)
    for i, doy in zip(np.where(left_mask)[0], doys[left_mask]):
        band = stack[i]
        cond = (band <= thresh_left) & (band > min_left)
        sos_doy[cond] = float(doy)

    eos_doy = np.full((H,W), np.nan, dtype=np.float32)
    for i, doy in zip(np.where(right_mask)[0], doys[right_mask]):
        band = stack[i]
        cond = (band <= thresh_right) & (band > min_right)
        eos_doy[cond] = float(doy)

    sowing  = replace_outliers(sos_doy - emergence)
    harvest = replace_outliers(eos_doy)
    duration= harvest - sowing
    duration= np.where(duration < 0, (365 - sowing) + harvest, duration)
    duration= replace_outliers(duration)
    return sowing, harvest, duration


def first_deriv_pixelwise(stack, doys, emergence):
    """
    R: compute_first_deriv + detect_phenodates_1D.

    SOS = time-step of steepest RISE on the left limb (argmax d1).
    EOS = time-step of steepest DECLINE on the right limb (argmin d1).

 
    """
    peak_val, median_peak, _, _ = _peak_and_limbs(stack, doys)
    H, W = peak_val.shape

    dt = np.diff(doys).astype(np.float32); dt[dt == 0] = 1.0
    d1_doys = ((doys[:-1] + doys[1:]) / 2.0).astype(np.float32)

    sos      = np.full((H, W), np.nan,      dtype=np.float32)
    sos_best = np.full((H, W), -np.inf,     dtype=np.float32)
    eos      = np.full((H, W), np.nan,      dtype=np.float32)
    eos_best = np.full((H, W),  np.inf,     dtype=np.float32)

    for i in range(len(d1_doys)):
        band = (stack[i+1].astype(np.float32) - stack[i].astype(np.float32)) / dt[i]
        valid = ~np.isnan(band)
        d = float(d1_doys[i])
        if d <= median_peak:          # left limb - track argmax
            better = valid & (band > sos_best)
            sos_best[better] = band[better]
            sos[better] = d
        else:                         # right limb - track argmin
            better = valid & (band < eos_best)
            eos_best[better] = band[better]
            eos[better] = d

    sos = np.where(np.isnan(peak_val), np.nan, sos).astype(np.float32)
    eos = np.where(np.isnan(peak_val), np.nan, eos).astype(np.float32)

    sowing   = replace_outliers(sos - emergence)
    harvest  = replace_outliers(eos)
    duration = harvest - sowing
    duration = np.where(duration < 0, (365 - sowing) + harvest, duration)
    duration = replace_outliers(duration)
    return sowing, harvest, duration


def _zero_cross_streaming(stack, doys, emergence, H, W, median_peak):
    """
    Zero-crossing detector for second_deriv_pixelwise, streaming over
    time steps to avoid materialising a full (T, H, W) d1/d2 array.

    Keeps only TWO consecutive d1 bands in memory at once, computes d2
    as the finite difference of those bands, and detects sign changes
    (negative→positive = local minimum in d1, i.e. inflection of the
    greening curve) on-the-fly.

    Left-limb rule: last crossing before median_peak → SOS.
    Right-limb rule: first crossing after median_peak → EOS.
    """
    dt = np.diff(doys).astype(np.float32); dt[dt == 0] = 1.0
    d1_doys = ((doys[:-1] + doys[1:]) / 2.0).astype(np.float32)

    sos       = np.full((H, W), np.nan,  dtype=np.float32)
    eos       = np.full((H, W), np.nan,  dtype=np.float32)
    eos_found = np.zeros((H, W), dtype=bool)

    prev_d1 = None
    prev_t  = None

    for i in range(len(d1_doys)):
        curr_d1 = (stack[i+1].astype(np.float32) - stack[i].astype(np.float32)) / dt[i]
        curr_t  = float(d1_doys[i])

        if prev_d1 is not None:
            # sign change from negative to positive = neg→pos crossing (local min of VI → SOS)
            crossing = ((np.sign(prev_d1) < 0) & (np.sign(curr_d1) >= 0) &
                        ~np.isnan(prev_d1) & ~np.isnan(curr_d1))
            d_cross = (prev_t + curr_t) / 2.0
            if curr_t <= median_peak:
                # left limb: keep overwriting so we get the LAST crossing
                sos[crossing] = d_cross
            else:
                # right limb: keep only the FIRST crossing
                update = crossing & ~eos_found
                eos[update] = d_cross
                eos_found[update] = True

        prev_d1 = curr_d1
        prev_t  = curr_t

    return sos, eos


def second_deriv_pixelwise(stack, doys, emergence):
    """
    R: compute_second_deriv + detect_phenodates_2D.

    Zero-crossings of d2 (sign change neg→pos) locate inflection points.
    Streaming implementation - only two consecutive d1 bands in memory
    at once, same memory motivation as first_deriv_pixelwise.
    """
    peak_val, median_peak, _, _ = _peak_and_limbs(stack, doys)
    H, W = peak_val.shape

    sos, eos = _zero_cross_streaming(stack, doys, emergence, H, W, median_peak)

    sos = np.where(np.isnan(peak_val), np.nan, sos).astype(np.float32)
    eos = np.where(np.isnan(peak_val), np.nan, eos).astype(np.float32)

    sowing   = replace_outliers(sos - emergence)
    harvest  = replace_outliers(eos)
    duration = harvest - sowing
    duration = np.where(duration < 0, (365 - sowing) + harvest, duration)
    duration = replace_outliers(duration)
    return sowing, harvest, duration


def sg_deriv_pixelwise(stack, doys, emergence, sg_window2=7, sg_poly2=3):
    """R: compute_SG_deriv + detect_phenodates_SGderiv - 2nd SG pass + Method A."""
    T, H, W = stack.shape
    flat = stack.reshape(T, -1)
    win = min(sg_window2, T); win = win if win % 2 == 1 else win - 1

    # savgol_filter rejects ANY NaN anywhere in its input (a global check,
    # not per-column like nanargmax) - masked-out pixels are all-NaN
    # columns here, so filter ONLY the fully-valid columns and leave the
    # rest untouched (still NaN).
    flat_sg2 = flat.copy()
    valid_cols = ~np.all(np.isnan(flat), axis=0)
    if win >= sg_poly2 + 2 and valid_cols.any():
        flat_sg2[:, valid_cols] = savgol_filter(flat[:, valid_cols], win, sg_poly2, axis=0)
    stack_sg2 = flat_sg2.reshape(T, H, W)
    return first_deriv_pixelwise(stack_sg2, doys, emergence)


# =============================================================================
# 11. VALIDATION HELPERS
# =============================================================================

def to_doy(date_str):
    """
    Convert a date-like value to day-of-year. Confirmed GT file format is
    day-first (e.g. 25-03-2020 = March 25) - NEVER falls back to pandas'
    default month-first guess, which would silently swap day/month for
    any ambiguous date (day ≤ 12) with no error or warning.
    """
    s = str(date_str).strip()
    for fmt in ("%d-%m-%Y", "%d/%m/%Y"):
        try:
            return datetime.strptime(s[:10], fmt).timetuple().tm_yday
        except Exception:
            continue
    # Fallback for already-parsed datetime/Timestamp objects, or minor
    # formatting variations (different separators, no leading zeros) -
    # still explicitly day-first, never the ambiguous default.
    try:
        return pd.to_datetime(date_str, dayfirst=True, errors="raise").dayofyear
    except Exception:
        return np.nan

def iqr_ok(series):
    q1, q3 = series.quantile([0.25, 0.75]); iqr = q3 - q1
    return (series > q1 - 1.5*iqr) & (series < q3 + 1.5*iqr)

def extract_raster_at_points(tif_path, lons, lats, nodata=-9999):
    with rasterio.open(tif_path) as src:
        vals = []
        for x, y in zip(lons, lats):
            try:
                v = list(src.sample([(x,y)]))[0][0]
                vals.append(np.nan if v == nodata else float(v))
            except Exception:
                vals.append(np.nan)
    return np.array(vals)

def bias_r_rmse(obs, est):
    ok = ~(np.isnan(obs) | np.isnan(est))
    if ok.sum() < 2: return np.nan, np.nan, np.nan
    bias = round((obs[ok]-est[ok]).mean(), 1)
    r, _ = pearsonr(obs[ok], est[ok])
    rmse = np.sqrt(((obs[ok]-est[ok])**2).mean())
    return bias, r, rmse

def categorise(series, q33=None, q66=None, labels=("Early","Mid","Late")):
    if q33 is None: q33 = series.quantile(0.33)
    if q66 is None: q66 = series.quantile(0.66)
    return pd.cut(series, bins=[-np.inf, q33, q66, np.inf], labels=labels)

def confusion_heatmap(ax, cm_arr, labels, title):
    im = ax.imshow(cm_arr[::-1], cmap="cividis", aspect="auto")
    ax.set_xticks(range(len(labels))); ax.set_xticklabels(labels, fontsize=8)
    ax.set_yticks(range(len(labels))); ax.set_yticklabels(labels[::-1], fontsize=8)
    for i in range(len(labels)):
        for j in range(len(labels)):
            ax.text(j, i, str(cm_arr[::-1][i,j]), ha="center", va="center", fontsize=8,
                    color="white" if cm_arr[::-1][i,j] < cm_arr.max()/2 else "black")
    ax.set_xlabel("Observed", fontsize=8); ax.set_ylabel("Estimated", fontsize=8)
    ax.set_title(title, fontsize=9)
    plt.colorbar(im, ax=ax, fraction=0.03)


# =============================================================================
# 12. MAPPING HELPERS
# =============================================================================

def _boundary_extent(ax, arr, aoi_gdf, transform):
    """Shared helper for doy_map/cat_map: draw the AOI/admin boundary as a
    solid grey background polygon, and return the (left,right,bottom,top)
    extent so the raster can be overlaid at its correct geographic position
    on top of it. Returns None if no boundary/transform was given, in which
    case the caller falls back to a plain pixel-grid imshow as before."""
    if aoi_gdf is None or transform is None:
        return None
    aoi_gdf.plot(ax=ax, color="#d9d9d9", edgecolor="#999999", linewidth=0.4, zorder=0)
    H, W = arr.shape
    left,  top    = transform * (0, 0)
    right, bottom = transform * (W, H)
    ax.set_xlim(left, right); ax.set_ylim(bottom, top)
    return (left, right, bottom, top)


def generate_validation_report(out_path, tag, index, country, useCaseName, crop,
                                Planting_year, method_label,
                                global_stats, dist_df, figure_paths,
                                thr=None):
    """
    Write a Word (.docx) validation report for one phenology run.

    Parameters
    ----------
    out_path     : Path - full path to write the .docx file
    tag          : str  - the run tag (embedded in all output filenames)
    index        : str  - index / combined index string used in this run
    country, useCaseName, crop, Planting_year : run metadata
    method_label : str  - e.g. "Threshold", "FirstDeriv", "CrossValidation"
    global_stats : dict - {"bias": float, "r": float, "rmse": float, "n": int}
                   Any value can be None if not available.
    dist_df      : pd.DataFrame or None - district-level metrics table
                   (columns: district, n, bias_days, r, rmse_days)
    figure_paths : list[Path] - PNGs to embed, in order; non-existent paths skipped
    thr          : tuple or None - (thr_sos, thr_eos) for Threshold runs

    Returns the Path written, or None if python-docx is not installed.
    """
    try:
        from docx import Document as DocxDocument
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import lxml.etree as etree
    except ImportError:
        print("  [Report] python-docx not installed - skipping report generation.")
        print("           Install with:  pip install python-docx")
        return None

    from datetime import datetime

    doc = DocxDocument()

    # ── Page margins (A4, 2cm all sides) ────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Helper: styled paragraph ─────────────────────────────────────────────
    def heading(text, level=1):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        return p

    def body(text):
        p = doc.add_paragraph(text)
        p.paragraph_format.space_after = Pt(4)
        return p

    def add_table_row(table, cells, bold=False, shade=None):
        row = table.add_row()
        for i, (cell, val) in enumerate(zip(row.cells, cells)):
            cell.text = str(val)
            if bold:
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            if shade:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), shade)
                tcPr.append(shd)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(2)

    # ── Title block ──────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(f"Phenology Validation Report")
    run.bold = True; run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.add_run(f"{country} | {useCaseName} | {crop} | Season {Planting_year}").font.size = Pt(12)

    doc.add_paragraph()  # spacer

    # ── Run metadata ─────────────────────────────────────────────────────────
    heading("1. Run Configuration", level=1)
    meta = [
        ("Country",        country),
        ("Use case",       useCaseName),
        ("Crop",           crop),
        ("Season year",    str(Planting_year)),
        ("Method",         method_label),
        ("Index / fusion", index),
    ]
    if thr is not None:
        meta.append(("Threshold (sos / eos)", f"{thr[0]:.2f} / {thr[1]:.2f}"))
    meta.append(("Run tag", tag))
    meta.append(("Report generated", datetime.now().strftime("%Y-%m-%d %H:%M")))

    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    tbl.columns[0].width = Inches(2.2)
    tbl.columns[1].width = Inches(4.5)
    for i, (k, v) in enumerate(meta):
        shade = "EBF3FB" if i % 2 == 0 else "FFFFFF"
        add_table_row(tbl, [k, v], shade=shade)

    doc.add_paragraph()

    # ── Global validation metrics ─────────────────────────────────────────────
    heading("2. Global Planting-Date Validation Metrics", level=1)
    body("Planting-only validation against ground-truth survey records "
        f"from {Planting_year}, after year-matching and IQR filtering.")
    doc.add_paragraph()

    gn    = global_stats.get("n")
    gbias = global_stats.get("bias")
    gr    = global_stats.get("r")
    grmse = global_stats.get("rmse")

    g_rows = [
        ("Metric", "Value"),
        ("GT points used (after IQR filter)", str(gn)   if gn    is not None else "-"),
        ("Bias (Obs − Est)",                  f"{gbias:.1f} days" if gbias is not None else "-"),
        ("Pearson R",                         f"{gr:.3f}"         if gr    is not None else "-"),
        ("RMSE",                              f"{grmse:.1f} days" if grmse is not None else "-"),
    ]
    g_tbl = doc.add_table(rows=0, cols=2)
    g_tbl.style = "Table Grid"
    g_tbl.columns[0].width = Inches(3.5)
    g_tbl.columns[1].width = Inches(3.2)
    for i, (k, v) in enumerate(g_rows):
        shade = "1F497D" if i == 0 else ("EBF3FB" if i % 2 == 1 else "FFFFFF")
        bold  = (i == 0)
        add_table_row(g_tbl, [k, v], bold=bold, shade=shade)
        if i == 0:
            for cell in g_tbl.rows[0].cells:
                for run in cell.paragraphs[0].runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    doc.add_paragraph()

    # ── District-level table ─────────────────────────────────────────────────
    if dist_df is not None and not dist_df.empty:
        heading("3. District-Level Validation", level=1)
        body("Metrics per sub-county/district, derived by spatial join of "
            "ground-truth points onto the AOI boundary polygons. "
            "Districts with fewer than 2 points are excluded. "
            "Sorted by RMSE (best first).")
        doc.add_paragraph()

        cols = ["District", "N", "Bias (days)", "Pearson R", "RMSE (days)"]
        col_keys = ["district", "n", "bias_days", "r", "rmse_days"]
        col_widths = [Inches(1.8), Inches(0.6), Inches(1.1), Inches(1.1), Inches(1.1)]

        d_tbl = doc.add_table(rows=0, cols=len(cols))
        d_tbl.style = "Table Grid"
        for i, w in enumerate(col_widths):
            d_tbl.columns[i].width = w

        # Header row
        hrow = d_tbl.add_row()
        for i, (cell, lbl) in enumerate(zip(hrow.cells, cols)):
            cell.text = lbl
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            tc = cell._tc; tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd"); shd.set(qn("w:val"),"clear")
            shd.set(qn("w:color"),"auto"); shd.set(qn("w:fill"),"1F497D")
            tcPr.append(shd)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(2)

        for ri, (_, row) in enumerate(dist_df.iterrows()):
            shade = "EBF3FB" if ri % 2 == 0 else "FFFFFF"
            vals = []
            for k in col_keys:
                v = row.get(k, "")
                if k in ("bias_days", "rmse_days") and v != "" and v == v:
                    vals.append(f"{float(v):.1f}")
                elif k == "r" and v != "" and v == v:
                    vals.append(f"{float(v):.3f}")
                else:
                    vals.append(str(v) if v == v else "-")
            add_table_row(d_tbl, vals, shade=shade)

        doc.add_paragraph()
        sect_offset = 1
    else:
        sect_offset = 0

    # ── Figures ───────────────────────────────────────────────────────────────
    fig_sect = 3 + sect_offset
    heading(f"{fig_sect}. Validation Figures", level=1)
    fig_labels = {
        "scatter":      "Scatter plot: Observed vs Estimated planting date (DOY), coloured by district.",
        "confusion":    "Category confusion matrix: Early / Mid / Late planting classification.",
        "district_bars":"District-level RMSE (left) and Bias (right) bar charts.",
    }
    any_fig = False
    for p_fig in figure_paths:
        p_fig = Path(p_fig)
        if not p_fig.exists():
            continue
        any_fig = True
        # Derive a caption from the filename
        stem = p_fig.stem.lower()
        caption = next((v for k, v in fig_labels.items() if k in stem),
                      p_fig.name)
        cap_p = doc.add_paragraph(caption)
        cap_p.runs[0].bold = True; cap_p.runs[0].font.size = Pt(9)
        cap_p.paragraph_format.space_after = Pt(3)
        try:
            doc.add_picture(str(p_fig), width=Inches(6.3))
        except Exception as e:
            body(f"[Could not embed figure: {p_fig.name} - {e}]")
        doc.add_paragraph()

    if not any_fig:
        body("No figure files found at the expected paths.")

    doc.save(str(out_path))
    return out_path


def doy_map(arr, title, save_path, aoi_gdf=None, transform=None):
    """
    Per-pixel DOY map. Pass aoi_gdf (from get_boundary()) and transform
    (from the stack's rasterio profile) to draw the admin/AOI boundary as a
    grey background and overlay only the in-mask pixels on top (NaN - i.e.
    outside the crop mask, or no valid data - rendered fully transparent),
    instead of a plain pixel grid with no reference geometry. Omit both to
    get the old plain-grid behaviour.

    interpolation="nearest" is forced on the imshow call: matplotlib's
    default antialiased/area-downsampling filter blends each pixel with
    its neighbors whenever the source array (often thousands of pixels per
    side for a province/country AOI) is squeezed into an ~8x6in figure.
    With sparse, salt-and-pepper crop-mask data that means almost every
    colored pixel partially blends into its transparent neighbors, showing
    up as a grey haze/halo over the whole map instead of crisp dots.
    "nearest" picks one source pixel per output pixel with no blending.
    """
    fig, ax = plt.subplots(figsize=(8,6))
    vmin, vmax = (np.nanpercentile(arr,[2,98]) if not np.all(np.isnan(arr)) else (0,1))
    extent = _boundary_extent(ax, arr, aoi_gdf, transform)

    cmap = plt.get_cmap("RdYlGn").copy()  # .copy() - never mutate the shared global cmap
    cmap.set_bad(alpha=0)
    im = ax.imshow(arr, cmap=cmap, vmin=vmin, vmax=vmax, extent=extent, zorder=1,
                   interpolation="nearest")
    plt.colorbar(im, ax=ax, label="DOY", fraction=0.03)
    ax.set_title(title, fontsize=10); ax.axis("off")
    plt.tight_layout(); plt.savefig(save_path, dpi=150, bbox_inches="tight"); plt.close()

def cat_map(arr, title, labels, colors, save_path, aoi_gdf=None, transform=None):
    """Categorical (Early/Mid/Late etc.) map - see doy_map() for the
    aoi_gdf/transform boundary-background behaviour."""
    q33, q66 = np.nanpercentile(arr, [33,66])
    cat = np.select([arr<q33, (arr>=q33)&(arr<q66), arr>=q66], [1,2,3])
    cat = np.where(np.isnan(arr), np.nan, cat)
    cmap = mcolors.ListedColormap(colors)
    cmap.set_bad(alpha=0)
    fig, ax = plt.subplots(figsize=(8,6))
    extent = _boundary_extent(ax, arr, aoi_gdf, transform)
    im = ax.imshow(cat, cmap=cmap, vmin=0.5, vmax=3.5, extent=extent, zorder=1,
                   interpolation="nearest")
    cb = plt.colorbar(im, ax=ax, fraction=0.03, ticks=[1,2,3])
    cb.ax.set_yticklabels(labels)
    ax.set_title(title, fontsize=10); ax.axis("off")
    plt.tight_layout(); plt.savefig(save_path, dpi=150, bbox_inches="tight"); plt.close()

def admin_map(arr, aoi_gdf, transform, district_col, title, save_path):
    stats = []
    for _, row in aoi_gdf.iterrows():
        m = geometry_mask([row.geometry.__geo_interface__], transform=transform,
                          invert=True, out_shape=arr.shape)
        v = arr[m]; v = v[~np.isnan(v)]
        stats.append({district_col: row[district_col],
                      "median": float(np.median(v)) if len(v) else np.nan,
                      "sd": float(np.std(v)) if len(v) else np.nan})
    merged = aoi_gdf.merge(pd.DataFrame(stats), on=district_col)
    fig, axes = plt.subplots(1, 2, figsize=(14,5))
    merged.plot(column="median", cmap="viridis", legend=True, ax=axes[0],
               missing_kwds={"color":"lightgrey"})
    axes[0].set_title(f"{title} - Median", fontsize=9); axes[0].axis("off")
    merged.plot(column="sd", cmap="magma", legend=True, ax=axes[1],
               missing_kwds={"color":"lightgrey"})
    axes[1].set_title(f"{title} - SD", fontsize=9); axes[1].axis("off")
    plt.tight_layout(); plt.savefig(save_path, dpi=150, bbox_inches="tight"); plt.close()


# =============================================================================
# 13. SAVE PHENOLOGY RASTER TRIO  (Sowing/Harvest/Duration)
# =============================================================================

def save_phenology_rasters(sowing, harvest, duration, profile, out_dir,
                           tag, method_label):
    out_prof = profile.copy()
    out_prof.update(count=1, dtype="float32", nodata=-9999)
    paths = {}
    for arr, metric in [(sowing,"PlantingDate"), (harvest,"HarvestDate"),
                        (duration,"CropDuration")]:
        p = out_dir / f"{tag}_{method_label}_{metric}.tif"
        data = np.where(np.isnan(arr), -9999, arr).astype(np.float32)
        with rasterio.open(p, "w", **out_prof) as dst: dst.write(data, 1)
        paths[metric] = p
    return paths


# =============================================================================
# 14. CHECKPOINT PLOTS RECONSTRUCTED FROM CACHE - NO RE-DOWNLOAD NEEDED
# =============================================================================
# Download_Stack_Smooth() caches every per-date composite in _tmp_download/
# and never deletes them, even after the smoothed stack is built. This lets
# any script reconstruct "raw vs smoothed" diagnostic plots straight from
# disk - no GEE calls, no re-download - by pairing those cached raw tiles
# with the already-saved smoothed stack.

def parse_stack_filename_years(stack_path):
    """Extract (Planting_year, Harvesting_year) from a stack filename like
    {country}_{useCaseName}_Sentinel_MultiIndex_{Py}_{Hy}_SG.tif"""
    stem = Path(stack_path).stem
    if stem.endswith("_SG"):
        stem = stem[:-3]
    nums = [p for p in stem.split("_") if p.isdigit() and len(p) == 4]
    if len(nums) >= 2:
        return int(nums[-2]), int(nums[-1])
    return None, None


def glob_cached_raw_tiles(base_dir, country, useCaseName, index, years):
    """
    Find cached per-date raw composite tiles for `index` across the given
    calendar year(s) - pass both Planting_year and Harvesting_year for a
    season that crosses a year boundary, so both halves are found.
    Returns sorted [(datetime, Path), ...].
    """
    tmp_dir = path_stack_dir(base_dir, country, useCaseName) / "_tmp_download"
    tiles, seen = [], set()
    for y in years:
        if y in seen or y is None: continue
        seen.add(y)
        for f in tmp_dir.glob(f"{index}_{y}*.tif"):
            datestr = f.stem.split("_")[-1]
            try:
                tiles.append((datetime.strptime(datestr, "%Y%m%d"), f))
            except ValueError:
                continue
    tiles.sort(key=lambda x: x[0])
    return tiles


def sample_checkpoint_examples_from_cache(base_dir, country, useCaseName, index,
                                          years, stack_sub, d_reg, transform,
                                          n_examples=3):
    """
    Reconstruct raw-vs-smoothed example pixel curves with NO GEE call:
      - stack_sub (T_reg,H,W): already-loaded SMOOTHED sub-stack for `index`
      - cached per-date raw GeoTIFFs in _tmp_download/ for the RAW side

    Pixel sampling uses real-world coordinates (not raw row/col), so this
    is correct even when the raw tile (e.g. an S1 band at its native scale)
    is on a different native grid than the final stack's reference grid.

    Returns a list of {"raw_dates","raw_vals","d_reg","sm_vals"} dicts.
    """
    from rasterio.transform import xy as rio_xy

    tiles = glob_cached_raw_tiles(base_dir, country, useCaseName, index, years)
    if not tiles:
        return []

    H, W = stack_sub.shape[1], stack_sub.shape[2]
    valid_frac = np.mean(~np.isnan(stack_sub), axis=0)
    ys, xs = np.where(valid_frac > 0.5)
    if len(ys) == 0:
        return []

    pick = np.linspace(0, len(ys) - 1, min(n_examples, len(ys))).astype(int)
    examples = []
    for p in pick:
        ry, rx = int(ys[p]), int(xs[p])
        lon, lat = rio_xy(transform, ry, rx)

        raw_dates, raw_vals = [], []
        for dt, tile_path in tiles:
            v = np.nan
            try:
                with rasterio.open(tile_path) as src:
                    sampled = list(src.sample([(lon, lat)]))[0][0]
                    if src.nodata is not None and sampled == src.nodata:
                        v = np.nan
                    elif abs(sampled) > 1e6:
                        v = np.nan
                    else:
                        v = float(sampled)
            except Exception:
                pass
            raw_dates.append(dt)
            raw_vals.append(v)

        examples.append({
            "raw_dates": raw_dates,
            "raw_vals": np.array(raw_vals, dtype=float),
            "d_reg": d_reg.copy(),
            "sm_vals": stack_sub[:, ry, rx].copy(),
        })
    return examples


def save_checkpoint_pdfs(checkpoint_data, plot_dir, country, useCaseName, year,
                         sg_window=7, sg_poly=3):
    """
    One PDF per index - raw (black, solid) vs SG-smoothed (red, dashed)
    for a few example pixels, faceted vertically. Matches the reference
    figure style: title, subtitle, 'Data' legend, real calendar dates.

    checkpoint_data : {index_name: {"examples": [...]}}
    Each example needs "raw_vals", "sm_vals", plus either "raw_dates"
    (real datetimes - preferred, robust to cross-year seasons) or
    "raw_doys" (converted via `year` as the reference year), and either
    "d_reg" or "sm_doys" for the smoothed-curve x-axis (DOY, converted
    via `year`).

    Returns list of saved PDF paths.
    """
    def doy_to_date(doy):
        return datetime(year, 1, 1) + timedelta(days=float(doy) - 1)

    saved_paths = []
    for name, data in checkpoint_data.items():
        examples = data.get("examples", [])
        if not examples:
            continue

        n = len(examples)
        fig, axes = plt.subplots(n, 1, figsize=(9, 2.6*n), sharex=True)
        if n == 1: axes = [axes]

        for i, (ax, ex) in enumerate(zip(axes, examples), start=1):
            raw_x = ex["raw_dates"] if "raw_dates" in ex else \
                    [doy_to_date(d) for d in ex["raw_doys"]]
            sm_key = "d_reg" if "d_reg" in ex else "sm_doys"
            sm_x = [doy_to_date(d) for d in ex[sm_key]]

            ax.plot(raw_x, ex["raw_vals"], color="black", lw=1.1,
                   marker="o", markersize=3, label="Raw")
            ax.plot(sm_x, ex["sm_vals"], color="red", lw=1.6,
                   linestyle="--", label="Smoothed")
            ax.set_ylabel(name, fontsize=10)
            ax.grid(alpha=0.2)
            ax.text(1.02, 0.5, f"Example {i}", transform=ax.transAxes,
                   rotation=270, va="center", fontsize=9)
            if i == 1:
                ax.legend(title="Data", loc="upper right", fontsize=8,
                         title_fontsize=9, frameon=False)

        axes[-1].set_xlabel("Date", fontsize=10)
        fig.suptitle(f"Smoothing noisy {name} time series with the\n"
                    f"Savitzky-Golay filter (window={sg_window}, poly={sg_poly})",
                    fontsize=13, x=0.02, ha="left", y=1.02)
        fig.text(0.02, 0.965, f"{country} {useCaseName}", fontsize=10, ha="left")
        plt.tight_layout(rect=[0, 0, 1, 0.93])

        p = plot_dir / f"checkpoint_{name}_{year}.pdf"
        plt.savefig(p, format="pdf", bbox_inches="tight")
        plt.close(fig)
        saved_paths.append(p)

    return saved_paths
